from docx import Document
import re
import os
from openpyxl import load_workbook, Workbook

def read_docx(file_path: str) -> str:
    """
    Extract the content of the DOCX file
    """
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        return '\n'.join(full_text)

    except Exception as e:
        return f"error: {e}"


def prepare_data(text: str) -> list[dict[str, str]]:
    """
    Find all phone numbers and names and at them to a list of dictionaries
    """
    data = []

    pattern = r'(?:[A-Za-z\s,]+)?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\s*,?\s*\+(\d{1,3}\s?\d{3}\s?\d{3}\s?\d{4})'

    matches = re.findall(pattern, text)

    for name, phone in matches:
        data.append({
            "name": name.strip(),
            "Phone Number": phone.strip()
        })

    return data


def open_excel(excel_path: str, data: list):
    """
    Open an Excel file and write data to it.
    """

    if os.path.exists(excel_path):
        workbook = load_workbook(excel_path)
        sheet = workbook.active
    else:
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Name", "Phone Number"])

    for item in data:
        sheet.append([item["name"], item["Phone Number"]])

    workbook.save(excel_path)


def remove_phone_numbers(input_file: str, output_file: str):
    """
    Remove all phone numbers from a DOCX file.
    """

    pattern = r'\+\d{1,3}\s?\d{3}\s?\d{3}\s?\d{4}'

    doc = Document(input_file)

    for para in doc.paragraphs:
        para.text = re.sub(pattern, "", para.text)

    doc.save(output_file)